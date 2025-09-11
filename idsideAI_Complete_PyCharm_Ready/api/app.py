import json
import os
import time

from docx import Document
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

app = FastAPI(title="IDECIDE Dummy Graph API", version="0.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


BASE = os.path.dirname(__file__)
DATA_FILE = os.path.join(BASE, "sample_graph.json")
with open(DATA_FILE) as f:
    SAMPLE = json.load(f)


def save_sample():
    with open(DATA_FILE, "w") as f:
        json.dump(SAMPLE, f, indent=2)


@app.get("/health")
def health():
    return {"status": "ok"}


@app.get("/graphs/{graph_id}/subgraph")
def subgraph(graph_id: str):
    edges = list(SAMPLE["edges"])  # copy
    # mark chosen on considers edge or add explicit
    ch = SAMPLE.get("chosen")
    if ch:
        for e in edges:
            if (
                e.get("type") in ("considers", "CHOSEN")
                and e.get("source") == ch.get("decision")
                and e.get("target") == ch.get("option")
            ):
                e["chosen"] = True
                e["type"] = e.get("type")
                break
        else:
            edges.append(
                {
                    "id": "CHOSEN",
                    "type": "CHOSEN",
                    "source": ch["decision"],
                    "target": ch["option"],
                    "chosen": True,
                }
            )
    return {"graph_id": graph_id, "nodes": SAMPLE["nodes"], "edges": edges}


@app.get("/graphs/{graph_id}/snapshots")
def snapshots(graph_id: str):
    return {"graph_id": graph_id, "snapshots": SAMPLE["snapshots"]}


@app.get("/graphs/{graph_id}/diff")
def diff(graph_id: str, frm: str = "v1", to: str = "v2"):
    return {"graph_id": graph_id, "from": frm, "to": to, "diff": SAMPLE["diff_v1_v2"]}


@app.post("/graphs/{graph_id}/choose")
def choose(graph_id: str, decision: str, option: str):
    SAMPLE["chosen"] = {"decision": decision, "option": option}
    save_sample()
    return {"ok": True, "chosen": SAMPLE["chosen"]}


@app.post("/graphs/{graph_id}/snapshot")
def create_snapshot(graph_id: str, root: str = None, depth: int = 3):
    # simulate by bumping a timestamp and maybe toggling a property
    now = SAMPLE.get("_now", 1) + 1
    SAMPLE["_now"] = now
    snap_id = f"S-{now}"
    # simulate: if chosen is OPT-A, add EV-99 if not present
    nodes = SAMPLE["nodes"]
    edges = SAMPLE["edges"]
    if not any(n.get("id") == "EV-99" for n in nodes):
        nodes.append(
            {
                "id": "EV-99",
                "type": "Evidence",
                "title": "SOC2 Report",
                "x": 340,
                "y": 430,
                "confidence": 0.6,
            }
        )
        edges.append(
            {"id": "E5", "type": "supported_by", "source": "EV-99", "target": "OPT-A"}
        )
    # snapshots list
    snaps = SAMPLE.setdefault("snapshots", [])
    snaps.append({"id": snap_id, "at": "2025-08-19T00:00:00Z"})
    save_sample()
    return {"snapshot": {"id": snap_id}}


from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
from starlette.staticfiles import StaticFiles

EXPORT_DIR = os.path.join(BASE, "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)
app.mount("/exports", StaticFiles(directory=EXPORT_DIR), name="exports")


def derive_why(sample):
    nodes = sample["nodes"]
    edges = sample["edges"]
    decision = next((n for n in nodes if n.get("type") == "Decision"), None)
    chosen = next(
        (
            e
            for e in edges
            if e.get("type", "").upper() == "CHOSEN" or e.get("chosen") == True
        ),
        None,
    )
    choice = next((n for n in nodes if chosen and n["id"] == chosen["target"]), None)
    supports = [
        e
        for e in edges
        if e.get("type") in ("supported_by", "SUPPORTED_BY")
        and e.get("target") == (choice["id"] if choice else None)
    ]
    evs = [n for n in nodes if n["id"] in [e["source"] for e in supports]]
    return {
        "title": f"Why: {choice['title'] if choice else 'Decision'}",
        "body": f"We selected {choice['title'] if choice else 'the option'} based on evidence and risk considerations.",
        "citations": [
            {
                "id": n["id"],
                "title": n.get("title", "Evidence"),
                "confidence": n.get("confidence"),
            }
            for n in evs
        ],
    }


@app.post("/graphs/{graph_id}/export/why")
def export_why(graph_id: str, decision: str = "", to: str = ""):
    why = derive_why(SAMPLE)
    pdf_name = f"why_{graph_id}_{int(time.time())}.pdf"
    pdf_path = os.path.join(EXPORT_DIR, pdf_name)
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    elements = []
    elements += [Paragraph("IDECIDE — Why Narrative", styles["Title"]), Spacer(1, 10)]
    elements += [
        Paragraph(why["title"], styles["Heading2"]),
        Paragraph(why["body"], styles["Normal"]),
        Spacer(1, 8),
    ]
    rows = [["ID", "Title", "Confidence"]] + [
        [c["id"], c["title"], str(c.get("confidence", ""))] for c in why["citations"]
    ]
    table = Table(rows, colWidths=[60, 240, 80])
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.grey)]))
    elements.append(table)
    doc.build(elements)

    docx_name = f"why_{graph_id}_{int(time.time())}.docx"
    docx_path = os.path.join(EXPORT_DIR, docx_name)
    d = Document()
    d.add_heading("IDECIDE — Why Narrative", 0)
    d.add_heading(why["title"], 1)
    d.add_paragraph(why["body"])
    for c in why["citations"]:
        d.add_paragraph(f"{c['id']}: {c['title']} (conf {c.get('confidence','')})")
    d.save(docx_path)
    return {"pdf": f"/exports/{pdf_name}", "docx": f"/exports/{docx_name}"}


@app.post("/graphs/{graph_id}/export/audit")
def export_audit(graph_id: str, frm: str = "v1", to: str = "v2"):
    fname = f"audit_{graph_id}_{frm}_{to}_{int(time.time())}.zip"
    fpath = os.path.join(EXPORT_DIR, fname)
    with zipfile.ZipFile(fpath, "w") as z:
        z.writestr(
            "manifest.json", json.dumps({"graph_id": graph_id, "from": frm, "to": to})
        )
        z.writestr(
            "snapshot_from.json",
            json.dumps({"nodes": SAMPLE["nodes"], "edges": SAMPLE["edges"]}),
        )
        z.writestr(
            "snapshot_to.json",
            json.dumps({"nodes": SAMPLE["nodes"], "edges": SAMPLE["edges"]}),
        )
    return {"url": f"/exports/{fname}"}


EXPORTS = os.path.join(BASE, "exports")
os.makedirs(EXPORTS, exist_ok=True)
app.mount("/exports", StaticFiles(directory=EXPORTS), name="exports")


def why_content(graph):
    # find chosen option
    chosen_opt = None
    for e in graph["edges"]:
        if e.get("type") in ("CHOSEN", "chosen") or e.get("chosen"):
            chosen_opt = e.get("target")
    if not chosen_opt:
        # pick best supported option
        support_counts = {}
        for e in graph["edges"]:
            if e["type"] == "supported_by":
                support_counts[e["target"]] = support_counts.get(e["target"], 0) + 1
        if support_counts:
            chosen_opt = max(support_counts, key=support_counts.get)
    # gather supporting evidence
    citations = []
    for e in graph["edges"]:
        if e["type"] == "supported_by" and e["target"] == chosen_opt:
            ev = next((n for n in graph["nodes"] if n["id"] == e["source"]), None)
            if ev:
                citations.append(ev)
    return chosen_opt, citations


@app.post("/graphs/{graph_id}/export/why")
def export_why(graph_id: str, decision: str = "DEC-1", to: str = "latest"):
    graph = SAMPLE
    chosen_opt, citations = why_content(graph)
    ts = time.strftime("%Y%m%d_%H%M%S")
    # PDF
    pdf_path = os.path.join(EXPORTS, f"Why_{ts}.pdf")
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    elems = [Paragraph("IDECIDE — Why Narrative", styles["Title"]), Spacer(1, 10)]
    elems.append(Paragraph(f"Decision: {decision}", styles["Heading2"]))
    elems.append(Paragraph(f"Chosen Option: {chosen_opt}", styles["Normal"]))
    elems.append(Spacer(1, 8))
    elems.append(Paragraph("Citations:", styles["Heading2"]))
    for ev in citations:
        line = f"{ev.get('title')} — {ev.get('source')} ({ev.get('date')}) — confidence {int((ev.get('confidence') or 0)*100)}%"
        if ev.get("url"):
            line += f" — {ev['url']}"
        elems.append(Paragraph(line, styles["Normal"]))
    doc.build(elems)
    # DOCX
    docx_path = os.path.join(EXPORTS, f"Why_{ts}.docx")
    d = Document()
    d.add_heading("IDECIDE — Why Narrative", 0)
    d.add_paragraph(f"Decision: {decision}")
    d.add_paragraph(f"Chosen Option: {chosen_opt}")
    d.add_heading("Citations", level=1)
    for ev in citations:
        line = f"{ev.get('title')} — {ev.get('source')} ({ev.get('date')}), conf {int((ev.get('confidence') or 0)*100)}%"
        if ev.get("url"):
            line += f" — {ev['url']}"
        d.add_paragraph(line)
    d.save(docx_path)
    return {
        "pdf": f"/exports/{os.path.basename(pdf_path)}",
        "docx": f"/exports/{os.path.basename(docx_path)}",
    }


@app.post("/graphs/{graph_id}/export/audit")
def export_audit(graph_id: str, frm: str = "v1", to: str = "v2"):
    ts = time.strftime("%Y%m%d_%H%M%S")
    zipname = f"audit_{ts}.zip"
    zpath = os.path.join(EXPORTS, zipname)
    import zipfile

    with zipfile.ZipFile(zpath, "w") as z:
        z.writestr(
            "manifest.json",
            json.dumps({"graph_id": graph_id, "from": frm, "to": to}, indent=2),
        )
        z.writestr("snapshot_from.json", json.dumps(SAMPLE, indent=2))
        z.writestr("snapshot_to.json", json.dumps(SAMPLE, indent=2))
    return {"url": f"/exports/{zipname}"}
